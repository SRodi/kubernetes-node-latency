"""Programmatic run-comparison report generator.

Replaces the prompt-based workflow (`docs/analysis-prompt.md`) with a CLI:

    python -m src.cli report --last 2
    python -m src.cli report 20260514-164726 20260514-164732

Produces a Markdown file and a matching .docx in `./analysis/` (repo root).
Each run's `phase_profile.png` is embedded at the top so the reader can
visually compare critical-path shape before reading the numeric tables.
"""
from __future__ import annotations

import datetime as _dt
import json
from collections import Counter
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

import pandas as pd

# Cilium configmap keys to diff across runs (subset of analysis-prompt.md).
# Provider-specific CNI/dataplane pods to keep when generating the report's
# focused phase_profile. Keeps the report visually consistent across
# providers by collapsing each plot to just the dataplane-owning pods.
REPORT_PODS_BY_PROVIDER: dict[str, set[str]] = {
    "gke_autopilot": {"anetd", "netd"},
    "gke_standard_dpv2": {"anetd", "netd"},
    "eks_eni_cilium": {"cilium"},
    "aks_overlay_cilium": {"cilium", "azure-cns"},
    "aks_byocni": {"cilium", "azure-cns"},
    "aks_kubenet": {"azure-cns"},
}

# Human-readable shorthand for each provider id. Used everywhere the
# report identifies a run (headline, KPI columns, cilium diff table) so
# readers see "AKS Overlay+Cilium" rather than `aks_overlay_cilium`.
PROVIDER_DISPLAY_NAMES: dict[str, str] = {
    "gke_autopilot":      "GKE Autopilot",
    "gke_standard_dpv2":  "GKE Standard (DPv2)",
    "aks_overlay_cilium": "AKS Overlay+Cilium",
    "aks_byocni":         "AKS BYO-CNI",
    "aks_kubenet":        "AKS Kubenet",
    "eks_eni_cilium":     "EKS ENI+Cilium",
}


def _provider_display(provider: str) -> str:
    return PROVIDER_DISPLAY_NAMES.get(provider, provider)


def _report_pods_filter(provider: str) -> set[str] | None:
    return REPORT_PODS_BY_PROVIDER.get(provider)


CILIUM_DIFF_KEYS = [
    "ipam", "routing-mode", "tunnel", "tunnel-protocol",
    "kube-proxy-replacement", "cni-chaining-mode",
    "identity-management-mode", "identity-allocation-mode",
    "enable-bandwidth-manager", "enable-bpf-masquerade",
    "enable-endpoint-routes", "enable-host-legacy-routing",
    "enable-ipv4-masquerade", "enable-l7-proxy",
    "prometheus-serve-addr", "operator-prometheus-serve-addr",
    "datapath-mode", "bpf-lb-mode", "bpf-lb-algorithm",
    "agent-not-ready-taint-key", "set-cilium-node-taints",
]

# Metrics to surface in the KPI table (order matters).
KPI_METRICS = [
    ("node_ready_after_register_s",   "node_ready_after_register_s (mean)  *(autoscaler-free)*"),
    ("node_register_latency_s",       "node_register_latency_s (mean)      *(autoscaler + VM bringup)*"),
    ("node_startup_latency_s",        "node_startup_latency_s (mean / p90)"),
    ("time_to_schedulable_s",         "time_to_schedulable_s (mean / p90)"),
    ("cni_conflist_install_s",        "cni_conflist_install_s (mean / p90)"),
    # T1\u2192T1c decomposition (from enrichment collector)
    ("pod_scheduling_lag_s",          "  \u21b3 pod scheduling lag (mean)"),
    ("image_pull_s",                  "  \u21b3 image pull (mean)"),
    ("csinode_block_s",               "  \u21b3 CSINode-ready offset (mean)"),
    ("taint_observed_offset_s",       "  \u21b3 cilium taint first-seen offset (mean)"),
    ("post_conflist_ready_s",         "post_conflist_ready_s (mean)"),
    ("cilium_init_duration_s",        "cilium_init_duration_s (mean)"),
    ("cni_induced_delay_s",           "cni_induced_delay_s (mean)"),
    ("cilium_scheduling_block_s",     "cilium_scheduling_block_s (mean / p90)"),
    ("cilium_bootstrap_total_s",      "bootstrap.overall (mean, range)"),
    ("cilium_bootstrap_early_init_s", "bootstrap.earlyInit (mean)"),
    ("cilium_bootstrap_k8s_init_s",   "bootstrap.k8sInit (mean)"),
    ("cilium_bootstrap_ipam_s",       "bootstrap.ipam (mean)"),
    ("cilium_bootstrap_maps_init_s",  "bootstrap.mapsInit (mean)"),
    ("cilium_bootstrap_bpf_base_s",   "bootstrap.bpfBase (mean)"),
    ("cilium_bootstrap_restore_s",    "bootstrap.restore (mean)"),
    ("cilium_endpoint_regen_avg_s",                "endpoint_regen avg (per-endpoint, mean)"),
    ("cilium_endpoint_regen_bpf_compilation_s",    "endpoint_regen bpfCompilation (mean)"),
    ("cilium_endpoint_regen_bpf_wait_for_elf_s",   "endpoint_regen bpfWaitForELF (mean)"),
    ("cilium_endpoint_regen_waiting_for_lock_s",   "endpoint_regen waitingForLock (mean)"),
]

# Anomaly event kinds worth surfacing (count > 0 → mention).
ANOMALY_KINDS = [
    "node_watch_reopen", "node_ready_watch_reopen",
    "cilium_deep_collect_failed",
    "node_schedulable_timeout",
    "uncordon_error",
]


@dataclass
class RunData:
    run_id: str
    run_dir: Path
    iterations: pd.DataFrame
    metadata: dict[str, Any]
    cilium_cm: dict[str, str] = field(default_factory=dict)
    cilium_summary: dict[str, Any] = field(default_factory=dict)
    event_counts: Counter = field(default_factory=Counter)
    missing: list[str] = field(default_factory=list)

    @property
    def provider(self) -> str:
        return (self.metadata.get("config") or {}).get("provider") or "unknown"

    @property
    def region(self) -> str:
        return (self.metadata.get("cluster") or {}).get("region") or "—"

    @property
    def k8s_version(self) -> str:
        return (self.metadata.get("cluster") or {}).get("kubernetes_version") or "—"

    @property
    def node_sku(self) -> str:
        nodes = (self.metadata.get("cluster") or {}).get("nodes") or []
        for n in nodes:
            labels = n.get("labels") or {}
            sku = labels.get("node.kubernetes.io/instance-type") \
                or labels.get("beta.kubernetes.io/instance-type")
            if sku:
                return sku
        return "—"

    @property
    def display_name(self) -> str:
        """Env shorthand (e.g. ``AKS Overlay+Cilium``) — used in tables."""
        return _provider_display(self.provider)

    @property
    def label(self) -> str:
        bits = [self.region]
        argv = self.metadata.get("cli_argv") or []
        if "--aks-node-provisioning" in argv:
            try:
                bits.append(argv[argv.index("--aks-node-provisioning") + 1])
            except IndexError:
                pass
        # Env-name-first label: "AKS Overlay+Cilium — 20260522-155452 (uksouth, nap)".
        return f"{self.display_name} — {self.run_id} ({', '.join(bits)})"

    @property
    def phase_profile_png(self) -> Path | None:
        """Provider-focused phase profile for the report.

        Generates a filtered variant (only the dataplane pods relevant to
        the provider — see ``REPORT_PODS_BY_PROVIDER``) on demand and
        falls back to the unfiltered ``phase_profile.png`` when filtering
        produces no plot or the provider has no mapping.
        """
        pods = _report_pods_filter(self.provider)
        unfiltered = self.run_dir / "plots" / "phase_profile.png"
        if not pods:
            return unfiltered if unfiltered.exists() else None
        suffix = "report"
        filtered = self.run_dir / "plots" / f"phase_profile_{suffix}.png"
        if not filtered.exists():
            try:
                from .plotting import _plot_phase_profile, _ok
                ok = _ok(self.iterations)
                if not ok.empty:
                    _plot_phase_profile(
                        ok,
                        self.run_dir / "plots",
                        title=f"({self.provider} @ {self.region})",
                        filename=f"phase_profile_{suffix}.png",
                        pods_filter=pods,
                    )
            except Exception:
                return unfiltered if unfiltered.exists() else None
        if filtered.exists():
            return filtered
        return unfiltered if unfiltered.exists() else None

    @property
    def phase_stacked_png(self) -> Path | None:
        p = self.run_dir / "plots" / "phase_stacked.png"
        return p if p.exists() else None

    @property
    def ok(self) -> pd.DataFrame:
        if "status" not in self.iterations.columns:
            return self.iterations
        return self.iterations[self.iterations["status"] == "success"]


def _augment_derived(df: pd.DataFrame) -> pd.DataFrame:
    """Backfill derived metrics for older iterations.csv files predating them."""
    if df.empty:
        return df
    def _delta(end: str, start: str) -> pd.Series:
        if end not in df.columns or start not in df.columns:
            return pd.Series([pd.NA] * len(df))
        return (pd.to_datetime(df[end], utc=True, errors="coerce")
                - pd.to_datetime(df[start], utc=True, errors="coerce")).dt.total_seconds().clip(lower=0)
    derived = {
        "node_ready_after_register_s": ("T4_node_ready", "T1_node_registered"),
        "cni_conflist_install_s":      ("T1c_cni_conflist", "T1_node_registered"),
        "post_conflist_ready_s":       ("T4_node_ready", "T1c_cni_conflist"),
        "cilium_scheduling_block_s":   ("T4b_schedulable", "T4_node_ready"),
        "time_to_schedulable_s":       ("T4b_schedulable", "T0_pod_created"),
    }
    for col, (end, start) in derived.items():
        if col not in df.columns or pd.to_numeric(df[col], errors="coerce").isna().all():
            s = _delta(end, start)
            if not s.isna().all():
                df[col] = s
    # Backfill trigger-pod lifecycle decomposition for legacy iterations.csv.
    try:
        from .analysis import enrich_trigger_pod_metrics
        enrich_trigger_pod_metrics(df)
    except Exception:
        pass
    return df


def load_run(run_dir: Path) -> RunData:
    rd = RunData(run_id=run_dir.name, run_dir=run_dir,
                 iterations=pd.DataFrame(), metadata={})

    csv = run_dir / "iterations.csv"
    if csv.exists():
        df = pd.read_csv(csv)
        # Backfill expanded cilium bootstrap sub-phase columns from any
        # iter-NNN/cilium_deep_headline.json payloads (older iterations.csv
        # files predate the schema; re-parse the source JSONs so reports get
        # the rich Cilium detail without rerunning the harness).
        from .plotting import _backfill_cilium_deep, _backfill_taint_observed
        _backfill_cilium_deep(df, run_dir)
        _backfill_taint_observed(df, run_dir)
        rd.iterations = _augment_derived(df)
    else:
        rd.missing.append("iterations.csv")

    md = run_dir / "run_metadata.json"
    if md.exists():
        try:
            rd.metadata = json.loads(md.read_text())
        except Exception:
            rd.missing.append("run_metadata.json (unparseable)")
    else:
        rd.missing.append("run_metadata.json")

    cm = run_dir / "cilium_config" / "cilium-config.json"
    if cm.exists():
        try:
            rd.cilium_cm = json.loads(cm.read_text()).get("data") or {}
        except Exception:
            rd.missing.append("cilium-config.json (unparseable)")
    else:
        rd.missing.append("cilium_config/cilium-config.json")

    cs = run_dir / "cilium_config" / "summary.json"
    if cs.exists():
        try:
            rd.cilium_summary = json.loads(cs.read_text())
        except Exception:
            pass

    re = run_dir / "raw_events.jsonl"
    if re.exists():
        with re.open() as f:
            for line in f:
                try:
                    rd.event_counts[json.loads(line).get("kind", "?")] += 1
                except Exception:
                    continue
    else:
        rd.missing.append("raw_events.jsonl")

    return rd


def resolve_runs(run_ids: list[str], *, last: int | None,
                  base_dir: Path) -> list[Path]:
    if run_ids:
        out: list[Path] = []
        for rid in run_ids:
            d = base_dir / rid
            if not d.is_dir():
                raise FileNotFoundError(f"no run dir: {d}")
            out.append(d)
        return out
    if last is None or last <= 0:
        raise ValueError("either run_ids or --last N must be provided")
    candidates = sorted(
        (p for p in base_dir.iterdir()
         if p.is_dir() and p.name not in {".gitkeep", "analysis"}
         and (p / "iterations.csv").exists()),
        key=lambda p: p.name,
        reverse=True,
    )
    if not candidates:
        raise FileNotFoundError(f"no run dirs with iterations.csv under {base_dir}")
    return list(reversed(candidates[:last]))


# ---------- KPI computation ----------

def _series(df: pd.DataFrame, col: str) -> pd.Series:
    if col not in df.columns:
        return pd.Series(dtype=float)
    return pd.to_numeric(df[col], errors="coerce").dropna()


def _fmt(val: float | None, digits: int = 2) -> str:
    if val is None or pd.isna(val):
        return "—"
    return f"{val:.{digits}f}"


def _kpi_cell(s: pd.Series, *, mode: str) -> str:
    if s.empty:
        return "—"
    if mode == "mean":
        return _fmt(s.mean())
    if mode == "mean_p90":
        return f"{_fmt(s.mean())} / {_fmt(s.quantile(0.9))}"
    if mode == "mean_range":
        return f"{_fmt(s.mean())} ({_fmt(s.min())} – {_fmt(s.max())})"
    return _fmt(s.mean())


def _kpi_row(metric: str, runs: list[RunData]) -> list[str]:
    mode = "mean"
    if metric in {"node_startup_latency_s", "time_to_schedulable_s",
                   "cni_conflist_install_s", "cilium_scheduling_block_s"}:
        mode = "mean_p90"
    if metric == "cilium_bootstrap_total_s":
        mode = "mean_range"
    return [_kpi_cell(_series(r.ok, metric), mode=mode) for r in runs]


def build_kpi_table(runs: list[RunData]) -> list[list[str]]:
    header = ["Metric"] + [r.label for r in runs]
    rows: list[list[str]] = [header]

    # static cluster-info rows
    rows.append(["environment"] + [r.display_name for r in runs])
    rows.append(["region"] + [r.region for r in runs])
    rows.append(["node SKU"] + [r.node_sku for r in runs])
    rows.append(["k8s version"] + [r.k8s_version for r in runs])
    rows.append(["iterations (ok / total)"] +
                [f"{(r.iterations['status'] == 'success').sum() if 'status' in r.iterations.columns else len(r.iterations)} / {len(r.iterations)}"
                 for r in runs])

    for col, label in KPI_METRICS:
        cells = _kpi_row(col, runs)
        if all(c == "—" for c in cells):
            continue
        rows.append([label] + cells)
    return rows


# ---------- Cilium config diff ----------

def _cilium_diff_data(runs: list[RunData]) -> tuple[list[list[str]], list[str], list[list[str]]]:
    """Structured form of the cilium config diff for docx rendering.

    Returns ``(diff_table_rows, identical_keys_lines, image_table_rows)``.
    Either table may be empty when no data is available.
    """
    diff_rows: list[list[str]] = []
    identical: list[str] = []
    image_rows: list[list[str]] = []
    if not any(r.cilium_cm for r in runs):
        return diff_rows, identical, image_rows

    diff_rows.append(["Key"] + [r.display_name for r in runs] + ["Status"])
    for key in CILIUM_DIFF_KEYS:
        vals = [r.cilium_cm.get(key) for r in runs]
        present = [(i, v) for i, v in enumerate(vals) if v is not None]
        if not present:
            continue
        cells = [v if v is not None else "—" for v in vals]
        if len(present) == 1:
            status = f"only in {runs[present[0][0]].display_name}"
        else:
            unique = {v for _, v in present}
            if len(unique) == 1 and len(present) == len(runs):
                identical.append(f"{key}={present[0][1]}")
                continue
            elif len(unique) == 1:
                status = "identical where present (missing elsewhere)"
            else:
                status = "differs"
        diff_rows.append([key] + cells + [status])
    if len(diff_rows) == 1:  # only header
        diff_rows = []

    image_runs: list[RunData] = []
    agent_imgs: list[str] = []
    operator_imgs: list[str] = []
    for r in runs:
        at = (r.cilium_summary.get("agent_template") or {}).get("containers") or []
        ot = (r.cilium_summary.get("operator_template") or {}).get("containers") or []
        if not at and not ot:
            continue
        agent = next((c.get("image", "—") for c in at if c.get("name") == "cilium-agent"), "—")
        operator = next((c.get("image", "—") for c in ot if c.get("name") == "cilium-operator"), "—")
        image_runs.append(r)
        agent_imgs.append(agent)
        operator_imgs.append(operator)
    if image_runs:
        image_rows.append(["Container"] + [r.display_name for r in image_runs])
        image_rows.append(["cilium-agent"] + agent_imgs)
        image_rows.append(["cilium-operator"] + operator_imgs)
    return diff_rows, identical, image_rows


def cilium_config_diff(runs: list[RunData]) -> list[str]:
    if not any(r.cilium_cm for r in runs):
        return ["_No `cilium_config/cilium-config.json` available for any run._"]

    # Build a table: rows = each diffable key; cols = each env.
    table_rows: list[list[str]] = []
    header = ["Key"] + [r.display_name for r in runs] + ["Status"]
    table_rows.append(header)

    identical_keys: list[str] = []
    for key in CILIUM_DIFF_KEYS:
        vals = [r.cilium_cm.get(key) for r in runs]
        present = [(i, v) for i, v in enumerate(vals) if v is not None]
        if not present:
            continue
        cells = [f"`{v}`" if v is not None else "—" for v in vals]
        if len(present) == 1:
            status = f"only in **{runs[present[0][0]].display_name}**"
        else:
            unique = {v for _, v in present}
            if len(unique) == 1 and len(present) == len(runs):
                identical_keys.append(f"`{key}`=`{present[0][1]}`")
                continue  # collapse identical rows into a single footer line
            elif len(unique) == 1:
                status = "identical where present (missing elsewhere)"
            else:
                status = "**differs**"
        table_rows.append([f"`{key}`"] + cells + [status])

    lines: list[str] = []
    if len(table_rows) > 1:
        lines.append(_table_md(table_rows))
    else:
        lines.append("_No differing keys across runs._")
    if identical_keys:
        lines.append("")
        lines.append("**Identical across all runs:** " + ", ".join(identical_keys))

    # Container images table
    img_table: list[list[str]] = []
    image_runs: list[RunData] = []
    agent_imgs: list[str] = []
    operator_imgs: list[str] = []
    for r in runs:
        at = (r.cilium_summary.get("agent_template") or {}).get("containers") or []
        ot = (r.cilium_summary.get("operator_template") or {}).get("containers") or []
        if not at and not ot:
            continue
        agent = next((c.get("image", "—") for c in at if c.get("name") == "cilium-agent"), "—")
        operator = next((c.get("image", "—") for c in ot if c.get("name") == "cilium-operator"), "—")
        image_runs.append(r)
        agent_imgs.append(agent)
        operator_imgs.append(operator)
    if image_runs:
        lines.append("")
        lines.append("**Container images:**")
        lines.append("")
        img_header = ["Container"] + [r.display_name for r in image_runs]
        img_table.append(img_header)
        img_table.append(["cilium-agent"] + [f"`{i}`" for i in agent_imgs])
        img_table.append(["cilium-operator"] + [f"`{i}`" for i in operator_imgs])
        lines.append(_table_md(img_table))

    return lines


# ---------- Anomalies ----------

def anomalies(runs: list[RunData]) -> list[str]:
    lines: list[str] = []
    for r in runs:
        bits: list[str] = []
        if r.iterations.empty:
            bits.append("no iterations.csv")
        else:
            n_total = len(r.iterations)
            n_ok = int((r.iterations.get("status") == "success").sum()) if "status" in r.iterations.columns else n_total
            if n_ok < n_total:
                bits.append(f"{n_total - n_ok} iteration(s) did not succeed")
        for kind in ANOMALY_KINDS:
            c = r.event_counts.get(kind, 0)
            if c:
                bits.append(f"{c}× `{kind}`")
        if bits:
            lines.append(f"- **{r.run_id}:** " + "; ".join(bits))
    return lines or ["_No anomalies detected._"]


# ---------- Auto-headline ----------

def headline(runs: list[RunData]) -> str:
    if len(runs) < 2:
        r = runs[0]
        for metric in ("node_ready_after_register_s", "node_startup_latency_s"):
            s = _series(r.ok, metric)
            if not s.empty:
                return (
                    f"Single-run report for **{r.display_name}** (`{r.run_id}`, "
                    f"{r.region}). "
                    f"`{metric}`: avg = {_fmt(s.mean())} s, "
                    f"p50 = {_fmt(s.median())} s, "
                    f"p90 = {_fmt(s.quantile(0.9))} s."
                )
        return (f"Single-run report for **{r.display_name}** "
                f"(`{r.run_id}`, {r.region}).")

    # Pick the best available headline metric (autoscaler-free preferred).
    metric = None
    for cand in ("node_ready_after_register_s", "node_startup_latency_s"):
        if all(not _series(r.ok, cand).empty for r in runs):
            metric = cand
            break
    if metric is None:
        return "Comparison report across the supplied runs."

    means = [_series(r.ok, metric).mean() for r in runs]
    fastest = min(range(len(runs)), key=lambda i: means[i] if pd.notna(means[i]) else float("inf"))
    slowest = max(range(len(runs)), key=lambda i: means[i] if pd.notna(means[i]) else -1.0)
    if fastest == slowest or pd.isna(means[fastest]) or pd.isna(means[slowest]):
        return "Comparison report across the supplied runs."
    delta = means[slowest] / means[fastest] if means[fastest] > 0 else float("inf")

    # Secondary: CNI conflist install delta if available
    cni = []
    for r in runs:
        s = _series(r.ok, "cni_conflist_install_s")
        cni.append(s.mean() if not s.empty else None)
    cni_note = ""
    if all(v is not None for v in cni) and cni[fastest] and cni[slowest]:
        cni_note = (f" CNI conflist install (`cni_conflist_install_s`, avg): "
                    f"{_fmt(cni[fastest])} s vs {_fmt(cni[slowest])} s.")

    qualifier = " (autoscaler-free)" if metric == "node_ready_after_register_s" else ""
    return (
        f"On `{metric}`{qualifier}, **{runs[fastest].display_name} is fastest** "
        f"(avg {_fmt(means[fastest])} s), vs **{runs[slowest].display_name}** "
        f"(avg {_fmt(means[slowest])} s, ~{delta:.1f}×).{cni_note} "
        f"See the per-run `phase_profile_report.png` plots for the phase decomposition."
    )


# ---------- Rendering ----------

def _table_md(rows: list[list[str]]) -> str:
    if not rows:
        return ""
    header, *body = rows
    width = len(header)
    out = ["| " + " | ".join(header) + " |"]
    out.append("|" + "|".join(["---"] * width) + "|")
    for row in body:
        out.append("| " + " | ".join(row) + " |")
    return "\n".join(out)


def render_markdown(runs: list[RunData], out_path: Path) -> Path:
    when = _dt.datetime.now(_dt.timezone.utc).isoformat(timespec="seconds")
    run_ids = ", ".join(r.run_id for r in runs)
    title = ("Run Analysis — " + " vs ".join(r.run_id for r in runs)
             if len(runs) > 1 else f"Run Report — {runs[0].run_id}")

    lines: list[str] = []
    lines.append(f"# {title}")
    lines.append("")
    lines.append(f"_Generated: {when}_  ")
    lines.append(f"_Runs analyzed: {run_ids}_")
    lines.append("")
    lines.append("## Phase profile")
    lines.append("")
    lines.append("Per-run mean phase Gantt (autoscaler/VM bringup shown as a number "
                 "in the subtitle so the post-T1 phases are readable).")
    lines.append("")
    for r in runs:
        p = r.phase_profile_png
        if p is None:
            lines.append(f"- _{r.run_id}: no `plots/phase_profile.png` (run with the harness ≥ T1c)._")
            continue
        # Markdown image path relative to the report file's parent directory:
        try:
            rel = Path("..") / p.resolve().relative_to(out_path.resolve().parent.parent)
        except ValueError:
            rel = p.resolve()
        lines.append(f"**{r.label}:**")
        lines.append("")
        lines.append(f"![phase_profile {r.run_id}]({rel.as_posix()})")
        lines.append("")

    lines.append("## Phase stacked")
    lines.append("")
    lines.append("Per-run stacked bar of the headline phases — quick visual "
                 "of how each phase contributes to total `time_to_schedulable`.")
    lines.append("")
    for r in runs:
        p = r.phase_stacked_png
        if p is None:
            lines.append(f"- _{r.run_id}: no `plots/phase_stacked.png`._")
            continue
        try:
            rel = Path("..") / p.resolve().relative_to(out_path.resolve().parent.parent)
        except ValueError:
            rel = p.resolve()
        lines.append(f"**{r.label}:**")
        lines.append("")
        lines.append(f"![phase_stacked {r.run_id}]({rel.as_posix()})")
        lines.append("")

    lines.append("## Headline")
    lines.append("")
    lines.append(headline(runs))
    lines.append("")

    # Cross-provider pod-running breakdown (≥ 2 runs only).
    pod_png = _build_pod_running_compare(runs, out_path.parent)
    if pod_png is not None:
        lines.append("## Pod-running breakdown (cross-provider)")
        lines.append("")
        lines.append(
            "Three panels.\n\n"
            "**(1) Two-lane Gantt** anchored at "
            "``T1_node_registered = 0``. Per provider:\n"
            "- **Top lane** = node bootstrap, split into two segments:\n"
            "  - **Green** = ``cni_conflist_install_s`` (``T1 → T1c``): "
            "kubelet finds a usable CNI conflist on disk; Node "
            "Ready=True follows almost immediately "
            "(``post_conflist_ready_s ≈ 0`` everywhere).\n"
            "  - **Amber** = ``cilium_scheduling_block_s`` "
            "(``T4 → T4b ≈ T1c → T4b``): the CNI agent's own "
            "post-conflist work *after* kubelet posts Ready — eBPF "
            "program load, identity allocation, endpoint sync, health "
            "checks — before it removes its own bootstrap NoSchedule "
            "taint (``node.cilium.io/agent-not-ready`` on Cilium, "
            "``node.cloudprovider.kubernetes.io/uninitialized`` on AKS "
            "byocni, etc.). **This is usually the largest unexplained "
            "gap** between network-ready and pod-schedulable.\n"
            "  - A dotted green tick marks ``T1c`` (network plugin "
            "installed) at the boundary between the two segments.\n"
            "- **Bottom lane (phase-colored)** = pod lifecycle "
            "starting at ``T_trigger_scheduled`` and decomposed into "
            "sandbox/CNI ADD → image pull → container create → "
            "container start.\n\n"
            "The *visual relationship* between the lanes is the story:\n"
            "- **AKS / EKS** — bottom lane starts at the **end of the "
            "amber segment** (not the end of the green one), because "
            "the bootstrap taint is the actual scheduler gate. Once "
            "scheduled, sandbox is quick (CNI was ready long before).\n"
            "- **GKE** — no blocking taint configured, so no amber "
            "segment at all. The bottom lane starts **before** the "
            "green tick — kube-scheduler binds the pod eagerly, and "
            "the sandbox segment extends past the tick (kubelet was "
            "holding sandbox creation until CNI conflist appeared). "
            "Same physical wait, just shifted from scheduler-time to "
            "kubelet-time.\n\n"
            "Each label is annotated with "
            "``taint_blocking_duration_s`` "
            "(= ``T4b_schedulable − T_taint_observed``) — present on "
            "AKS/EKS where the bootstrap taint exists, absent on GKE.\n\n"
            "**(2) Kubelet vs network split** inside the pod→running "
            "window — kubelet (prepull + create + run_gap) vs network "
            "(image pull) with the percentage annotated per provider. "
            "Surfaces whether kubelet wiring or registry pull is the "
            "bottleneck (e.g. on GKE the sandbox-absorbs-network-wait "
            "pattern + image pull pushes the network share to "
            "majority).\n\n"
            "**(3) CDF** of ``trigger_total_s`` (pod scheduled → "
            "Running) per provider — per-iteration spread."
        )
        lines.append("")
        try:
            rel = pod_png.resolve().relative_to(out_path.resolve().parent)
            rel_str = rel.as_posix()
        except ValueError:
            rel_str = pod_png.resolve().as_posix()
        lines.append(f"![compare_pod_running]({rel_str})")
        lines.append("")

    lines.append("## KPI table")
    lines.append("")
    lines.append(_table_md(build_kpi_table(runs)))
    lines.append("")

    lines.append("## Cilium config diff")
    lines.append("")
    lines.extend(cilium_config_diff(runs))
    lines.append("")

    lines.append("## Anomalies")
    lines.append("")
    lines.extend(anomalies(runs))
    lines.append("")

    caveats: list[str] = []
    for r in runs:
        for m in r.missing:
            caveats.append(f"- `{r.run_id}`: missing `{m}` — affected sections may be partial.")
    if caveats:
        lines.append("## Caveats")
        lines.append("")
        lines.extend(caveats)
        lines.append("")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text("\n".join(lines))
    return out_path


def render_docx(runs: list[RunData], out_path: Path) -> Path:
    from docx import Document
    from docx.shared import Inches, Pt

    doc = Document()
    when = _dt.datetime.now(_dt.timezone.utc).isoformat(timespec="seconds")
    title = ("Run Analysis — " + " vs ".join(r.run_id for r in runs)
             if len(runs) > 1 else f"Run Report — {runs[0].run_id}")
    doc.add_heading(title, level=0)
    doc.add_paragraph(f"Generated: {when}").italic = True
    doc.add_paragraph("Runs analyzed: " + ", ".join(r.run_id for r in runs)).italic = True

    # Phase profile images first
    doc.add_heading("Phase profile", level=1)
    doc.add_paragraph(
        "Per-run mean phase Gantt. The T0→T1 autoscaler/VM bringup is shown as "
        "a number in each plot's subtitle so the post-T1 lifecycle is readable."
    )
    for r in runs:
        doc.add_heading(r.label, level=2)
        p = r.phase_profile_png
        if p is None:
            doc.add_paragraph(f"(no phase_profile.png for {r.run_id})")
        else:
            doc.add_picture(str(p), width=Inches(6.5))

    doc.add_heading("Phase stacked", level=1)
    doc.add_paragraph(
        "Per-run stacked bar of the headline phases — quick visual of how "
        "each phase contributes to total time_to_schedulable."
    )
    for r in runs:
        doc.add_heading(r.label, level=2)
        p = r.phase_stacked_png
        if p is None:
            doc.add_paragraph(f"(no phase_stacked.png for {r.run_id})")
        else:
            doc.add_picture(str(p), width=Inches(6.5))

    doc.add_heading("Headline", level=1)
    doc.add_paragraph(headline(runs))

    pod_png = _build_pod_running_compare(runs, out_path.parent)
    if pod_png is not None:
        doc.add_heading("Pod-running breakdown (cross-provider)", level=1)
        doc.add_paragraph(
            "Three panels. (1) Two-lane Gantt anchored at "
            "T1_node_registered=0. TOP LANE = node bootstrap split "
            "into two segments: GREEN = cni_conflist_install_s "
            "(T1→T1c) — kubelet finds a usable CNI conflist on disk; "
            "AMBER = cilium_scheduling_block_s (T4→T4b ≈ T1c→T4b) — "
            "the CNI agent's own post-conflist work (eBPF / IPAM / "
            "identity / health) AFTER kubelet posts Ready=True, "
            "before it removes its bootstrap NoSchedule taint. "
            "Dotted green tick marks T1c (network plugin installed). "
            "BOTTOM LANE = pod lifecycle from T_trigger_scheduled "
            "decomposed into sandbox/CNI ADD → image pull → container "
            "create → container start. The visual relationship is the "
            "story: AKS/EKS bottom lane starts at the end of the "
            "AMBER segment (taint is the scheduler gate), then "
            "sandbox is quick; GKE has no amber segment (no blocking "
            "taint) and the bottom lane starts BEFORE the green tick "
            "— kubelet absorbs the network wait inside sandbox setup. "
            "Each label annotated with taint_blocking_duration_s "
            "(T4b - T_taint_observed). (2) Kubelet vs network split "
            "inside the pod→running window — kubelet (prepull + "
            "create + start) vs network (image pull) with percentage "
            "annotated. (3) CDF of trigger_total_s per provider — "
            "per-iteration spread."
        )
        try:
            doc.add_picture(str(pod_png), width=Inches(6.5))
        except Exception:
            doc.add_paragraph(f"(failed to embed {pod_png.name})")

    doc.add_heading("KPI table", level=1)
    rows = build_kpi_table(runs)
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Light Grid Accent 1"
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            table.rows[i].cells[j].text = cell
            for run in table.rows[i].cells[j].paragraphs[0].runs:
                run.font.size = Pt(9)
                if i == 0:
                    run.bold = True

    doc.add_heading("Cilium config diff", level=1)
    diff_rows, identical, image_rows = _cilium_diff_data(runs)
    if not diff_rows and not identical and not image_rows:
        doc.add_paragraph("No `cilium_config/cilium-config.json` available for any run.")
    else:
        if diff_rows:
            t = doc.add_table(rows=len(diff_rows), cols=len(diff_rows[0]))
            t.style = "Light Grid Accent 1"
            for i, row in enumerate(diff_rows):
                for j, cell in enumerate(row):
                    t.rows[i].cells[j].text = cell
                    for run in t.rows[i].cells[j].paragraphs[0].runs:
                        run.font.size = Pt(9)
                        if i == 0:
                            run.bold = True
        if identical:
            p = doc.add_paragraph()
            p.add_run("Identical across all runs: ").bold = True
            p.add_run(", ".join(identical))
        if image_rows:
            doc.add_paragraph().add_run("Container images:").bold = True
            t = doc.add_table(rows=len(image_rows), cols=len(image_rows[0]))
            t.style = "Light Grid Accent 1"
            for i, row in enumerate(image_rows):
                for j, cell in enumerate(row):
                    t.rows[i].cells[j].text = cell
                    for run in t.rows[i].cells[j].paragraphs[0].runs:
                        run.font.size = Pt(8)
                        if i == 0:
                            run.bold = True

    doc.add_heading("Anomalies", level=1)
    for line in anomalies(runs):
        if line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            doc.add_paragraph(line)

    caveats: list[str] = []
    for r in runs:
        for m in r.missing:
            caveats.append(f"{r.run_id}: missing {m}")
    if caveats:
        doc.add_heading("Caveats", level=1)
        for c in caveats:
            doc.add_paragraph(c, style="List Bullet")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


def output_basename(runs: list[RunData]) -> str:
    if len(runs) == 1:
        return f"report-{runs[0].run_id}"
    if len(runs) == 2:
        return f"compare-{runs[0].run_id}-vs-{runs[1].run_id}"
    return f"compare-{runs[0].run_id}-plus{len(runs)-1}"


def _build_pod_running_compare(runs: list[RunData], out_dir: Path) -> Path | None:
    """Generate the cross-provider pod-running breakdown into ``out_dir``.

    Reuses :func:`plotting._plot_compare_pod_running` so the report is
    self-contained (no dependency on ``analysis/`` having been refreshed).
    Returns the PNG path on success, ``None`` if fewer than 2 runs or
    the plot could not be produced.
    """
    if len(runs) < 2:
        return None
    csvs = [r.run_dir / "iterations.csv" for r in runs
            if (r.run_dir / "iterations.csv").exists()]
    if len(csvs) < 2:
        return None
    try:
        from .plotting import _plot_compare_pod_running
        return _plot_compare_pod_running(csvs, out_dir)
    except Exception:
        return None


def build_report(run_ids: list[str], *, last: int | None,
                 results_dir: Path, out_dir: Path) -> tuple[Path, Path]:
    """Resolve runs, build both .md and .docx in out_dir, return paths."""
    run_dirs = resolve_runs(run_ids, last=last, base_dir=results_dir)
    runs = [load_run(d) for d in run_dirs]
    base = output_basename(runs)
    md_path = render_markdown(runs, out_dir / f"{base}.md")
    docx_path = render_docx(runs, out_dir / f"{base}.docx")
    return md_path, docx_path
